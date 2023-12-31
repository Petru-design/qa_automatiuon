import os

import docx

from stiEF_tests.tests.utils.comparators import recursive_attribute_compare, compare_texts


font_comparable_fields = {
    "all_caps": type(None),
    "bold": bool,
    "color": {
        "rgb": type(None),
        "theme_color": type(None),
        "type": int,
    },
    "complex_script": type(None),
    "cs_bold": type(None),
    "cs_italic": type(None),
    "double_strike": type(None),
    "emboss": type(None),
    "hidden": type(None),
    "highlight_color": type(None),
    "imprint": type(None),
    "italic": type(None),
    "math": type(None),
    "name": type(None),
    "no_proof": type(None),
    "outline": type(None),
    "rtl": type(None),
    "shadow": type(None),
    "size": type(None),
    "small_caps": type(None),
    "snap_to_grid": type(None),
    "spec_vanish": type(None),
    "strike": type(None),
    "subscript": type(None),
    "superscript": type(None),
    "underline": type(None),
    "web_hidden": type(None),
}

paragraph_comparable_fields = {
    "alignment": type(None),
    "paragraph_format": {
        "alignment": type(None),
        "first_line_indent": type(None),
        "keep_together": type(None),
        "keep_with_next": type(None),
        "left_indent": type(None),
        "line_spacing": float,
        "line_spacing_rule": int,
        "page_break_before": bool,
        "right_indent": type(None),
        "space_after": docx.shared.Twips,
        "space_before": docx.shared.Twips,
    },
    "style": {
        "base_style": type(None),
        "builtin": bool,
        "font": font_comparable_fields,
        "hidden": bool,
        "locked": bool,
        "name": str,
        "paragraph_format": {
            "alignment": type(None),
            "first_line_indent": type(None),
            "keep_together": type(None),
            "keep_with_next": type(None),
            "left_indent": type(None),
            "line_spacing": docx.shared.Twips,
            "line_spacing_rule": int,
            "page_break_before": bool,
            "right_indent": type(None),
            "space_after": docx.shared.Twips,
            "space_before": docx.shared.Twips,
        },
        "priority": int,
        "quick_style": bool,
        "style_id": str,
        "type": int,
        "unhide_when_used": bool,
    },
}

run_comparable_fields = {
    "bold": bool,
    "font": font_comparable_fields,
    "italic": type(None),
    "style": {
        "base_style": type(None),
        "builtin": bool,
        "font": font_comparable_fields,
        "hidden": bool,
        "locked": bool,
        "name": str,
        "priority": int,
        "quick_style": bool,
        "style_id": str,
        "type": int,
        "unhide_when_used": bool,
    },
    "underline": type(None),
}


def test_docx_text(reference_path, subject_path, result_path, naming_prefix):
    baseline = docx.Document(reference_path)
    subject = docx.Document(subject_path)
    compare_texts(
        "\n".join(paragraph.text for paragraph in baseline.paragraphs),
        "\n".join(paragraph.text for paragraph in subject.paragraphs),
        os.path.join(result_path, f"{naming_prefix}_text_result.txt"),
        naming_prefix=naming_prefix,
    )


def record_result(ent_1, ent_2, result_path, exception):
    with open(result_path, "w", encoding="utf-8") as result:
        result.write(
            "\n".join((ent_1.text, "-" * 60, ent_2.text,
                      "-" * 60, str(exception)))
        )


def test_docx_format(reference_path, subject_path, result_path, naming_prefix):
    baseline = docx.Document(reference_path)
    subject = docx.Document(subject_path)
    result_path = os.path.join(result_path, "format_result.txt")

    for i, (baseline_paragraph, subject_paragraph) in enumerate(
        zip(baseline.paragraphs, subject.paragraphs)
    ):
        try:
            recursive_attribute_compare(
                baseline_paragraph, subject_paragraph, paragraph_comparable_fields
            )
        except AssertionError as e:
            record_result(baseline_paragraph,
                          subject_paragraph, result_path, e)
            raise Exception(
                f"Different formating in paragraph {
                    i}. See output at {result_path}"
            ) from e

        if len(baseline_paragraph.runs) != len(subject_paragraph.runs):
            exception = Exception(
                f"Different formating in paragraph {
                    i}. See output at {result_path}"
            )
            record_result(baseline_paragraph, subject_paragraph,
                          result_path, exception)
            raise exception

        for j, (baseline_run, subject_run) in enumerate(
            zip(baseline_paragraph.runs, subject_paragraph.runs)
        ):
            try:
                recursive_attribute_compare(
                    baseline_run, subject_run, run_comparable_fields
                )
            except AssertionError as e:
                record_result(baseline_run, subject_run, result_path, e)
                raise Exception(
                    f"Different formating in run {
                        j}. See output at {result_path}"
                ) from e
